import cv2
import tkinter as tk
from tkinter import messagebox
from datetime import datetime
from PIL import Image, ImageTk
import tensorflow as tf
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from tensorflow.keras.preprocessing import image
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# Cargar modelo base
modelo_base = tf.keras.applications.MobileNetV2(
    input_shape=(224, 224, 3),
    include_top=False,
    pooling='avg',
    weights='imagenet'
)

# Cargar embeddings y nombres
datos = np.load('embeddings.npz')
embeddings_guardados = datos['embeddings']
nombres_propietarios = datos['nombres']

def obtener_embedding(ruta):
    img = image.load_img(ruta, target_size=(224, 224))
    img_array = image.img_to_array(img) / 255.0
    img_array = np.expand_dims(img_array, axis=0)
    embedding = modelo_base.predict(img_array, verbose=0)
    return embedding[0]

def name_format(nombre):
    return nombre.replace("_", " ").title()

# Crear ventana principal
ventana = tk.Tk()
ventana.title("Identificador por Cámara - Sistema de Control")
ventana.geometry("1000x500")
ventana.configure(bg="white")

verde_crema = "#eef5e6"
verde_olivo = "#556b2f"

font_titulo = ("Times New Roman", 18, "bold")
font_italic = ("Times New Roman", 18, "italic")
font_normal = ("Times New Roman", 12)

# Encabezado
tk.Label(ventana, text="Identificador de Fierros", font=font_titulo, fg=verde_olivo, bg="white").pack()
tk.Label(ventana, text="Sistema de Control", font=font_titulo, fg=verde_olivo, bg="white").pack()
frame_header = tk.Frame(ventana, bg="white")
frame_header.pack()
tk.Label(frame_header, text="Villa de San Antonio, Comayagua", font=font_italic, fg=verde_olivo, bg="white").pack(side="left")

# Marco principal dividido horizontalmente
main_frame = tk.Frame(ventana, bg="white")
main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

# Columna izquierda: cámara y botones
left_frame = tk.Frame(main_frame, bg=verde_crema, padx=30, pady=20)
left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Vista previa de cámara
panel = tk.Label(left_frame)
panel.pack(pady=0)

# Botón de captura
boton_frame = tk.Frame(left_frame, bg=verde_crema)
boton_frame.pack(pady=5)

btn_capturar = tk.Button(boton_frame, text="📸 Capturar y Predecir", bg=verde_olivo, fg="white", font=("Times New Roman", 12, "bold"), command=lambda: capturar_y_predecir())
btn_capturar.pack(side=tk.LEFT, padx=10)

btn_guardar = tk.Button(boton_frame, text="💾 Guardar Resultado en Word", bg=verde_olivo, fg="white", font=("Times New Roman", 12, "bold"), command=lambda: guardar_en_word())
btn_reiniciar = tk.Button(boton_frame, text="🔁 Hacer Otra Predicción", bg=verde_olivo, fg="white", font=("Times New Roman", 12, "bold"), command=lambda: reiniciar_prediccion())

# Columna derecha: resultados
right_frame = tk.Frame(main_frame, bg=verde_crema, padx=10)
right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=False)

resultado_title = tk.Label(right_frame, text="Resultados                      ", font=font_titulo, bg=verde_crema, fg=verde_olivo)
resultado_title.pack(anchor="w")

resultado_lbl = tk.Label(right_frame, text="", font=font_normal, bg=verde_crema, wraplength=300, justify="left", padx=10, pady=10)
resultado_lbl.pack(pady=5, anchor="w")

# Footer
tk.Label(ventana, text="versión beta IS-702 04.25", font=("Times New Roman", 9, "italic"), fg="#4f4f4f", bg="white").pack(pady=8)

ruta_prediccion = ""
resultado_prediccion = ""

def mostrar_botones_secundarios():
    btn_guardar.pack(side=tk.LEFT, padx=10)
    btn_reiniciar.pack(side=tk.LEFT, padx=10)

def capturar_y_predecir():
    global ruta_prediccion, resultado_prediccion
    if current_frame is not None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_prediccion = f"Captura_{timestamp}.jpg"
        cv2.imwrite(ruta_prediccion, cv2.cvtColor(current_frame, cv2.COLOR_RGB2BGR))

        try:
            nuevo_embedding = obtener_embedding(ruta_prediccion)
            similitudes = cosine_similarity([nuevo_embedding], embeddings_guardados)[0]
            top3 = np.argsort(similitudes)[::-1][:3]

            resultado_prediccion = f"Más similar:\n{name_format(nombres_propietarios[top3[0]])} ({similitudes[top3[0]]:.4f})\n\n📝 Otros posibles:\n" + \
                f"• {name_format(nombres_propietarios[top3[1]])} ({similitudes[top3[1]]:.4f})\n" + \
                f"• {name_format(nombres_propietarios[top3[2]])} ({similitudes[top3[2]]:.4f})"
            resultado_lbl.config(text=resultado_prediccion)
            mostrar_botones_secundarios()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo identificar el fierro:\n{e}")

def guardar_en_word():
    if not ruta_prediccion or not resultado_prediccion:
        messagebox.showwarning("Advertencia", "Primero realiza una predicción.")
        return
    try:
        doc = Document()
        titulo = doc.add_heading("Resultado de Predicción de Fierro", 0)
        run = titulo.runs[0]
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_picture(ruta_prediccion, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # Centrado
        doc.add_paragraph(resultado_prediccion)
        nombre_doc = f"Reporte_prediccion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(nombre_doc)
        messagebox.showinfo("Guardado", f"El resultado se ha guardado en {nombre_doc}")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar el reporte:\n{e}")

def reiniciar_prediccion():
    resultado_lbl.config(text="")
    global ruta_prediccion, resultado_prediccion
    ruta_prediccion = ""
    resultado_prediccion = ""
    btn_guardar.pack_forget()
    btn_reiniciar.pack_forget()

def actualizar_video():
    global current_frame
    ret, frame = cap.read()
    if ret:
        frame = cv2.flip(frame, 1)
        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        current_frame = rgb
        img = Image.fromarray(rgb)
        imgtk = ImageTk.PhotoImage(image=img.resize((400, 250)))
        panel.imgtk = imgtk
        panel.config(image=imgtk)
    panel.after(10, actualizar_video)

global current_frame
current_frame = None
cap = cv2.VideoCapture(0)
actualizar_video()

ventana.bind("<space>", lambda event: capturar_y_predecir())
ventana.bind("<Escape>", lambda event: ventana.destroy())
ventana.mainloop()
cap.release()
cv2.destroyAllWindows()
